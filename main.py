from docx import Document


def bionify(path_to_text: str) -> None:
    doc = Document(path_to_text)
    new_doc = Document()
    all_paragraphs = doc.paragraphs
    for paragraph in all_paragraphs:
        word_list = paragraph.text.split(' ')
        new_paragraph = new_doc.add_paragraph()
        for word in word_list:
            i = 0
            while i < len(word):
                if i == 0 or i == 1:
                    new_paragraph.add_run(word[i]).bold = True
                else:
                    new_paragraph.add_run(word[i]).bold = False
                i += 1
            new_paragraph.add_run(' ')
    # Input the path to the document that you wish to save to:
    new_doc.save('sample_output.docx')


if __name__ == '__main__':
    # Input the path to the document containing your text file you wish to read from:
    bionify(r'C:\Users\Wrenpo\Desktop\BionicTexterizer\sample_input.docx')
